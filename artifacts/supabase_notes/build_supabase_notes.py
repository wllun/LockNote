from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Supabase.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5E6B7A"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F4F7FB"
PALE_GOLD = "FFF7E6"
GOLD = "9A6A00"
PALE_RED = "FFF2F2"
RED = "9B1C1C"
WHITE = "FFFFFF"
MONO = "Consolas"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=180, bottom=120, end=180) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_keep_with_next(paragraph, value=True) -> None:
    paragraph.paragraph_format.keep_with_next = value


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color: str, side="left", size=12, space=8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_hyperlink(paragraph, text: str, url: str, color=BLUE) -> None:
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr, fld_char_2])


def add_numbering_definition(doc: Document, abstract_id: int, num_id: int, fmt: str, text: str,
                             left=540, hanging=270) -> None:
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), fmt)
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), text)
    lvl.append(lvl_text)
    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), str(left))
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def add_bullet(doc, text: str, num_id=43):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, num_id)
    p.add_run(text)
    return p


def add_step(doc, lead: str, detail: str, num_id=44):
    p = doc.add_paragraph(style="Normal")
    apply_num(p, num_id)
    r = p.add_run(lead)
    r.bold = True
    p.add_run(detail)
    return p


def add_code(doc, lines: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    set_paragraph_shading(p, "F2F4F7")
    set_paragraph_border(p, BLUE, side="left", size=14, space=8)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, MONO, 9.2, NAVY)
    return p


def add_callout(doc, label: str, text: str, fill=PALE_BLUE, accent=BLUE, ink=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(9)
    p.paragraph_format.line_spacing = 1.18
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, accent, side="left", size=18, space=9)
    run = p.add_run(f"{label}  ")
    set_run_font(run, size=10.5, color=accent, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=ink)
    return p


def add_definition(doc, term: str, definition: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    term_run = p.add_run(f"{term}: ")
    set_run_font(term_run, size=11, color=NAVY, bold=True)
    p.add_run(definition)
    return p


def add_check(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    box = p.add_run("[ ]  ")
    set_run_font(box, MONO, 10.5, BLUE, bold=True)
    p.add_run(text)
    return p


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in specs.items():
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def add_header_footer(section) -> None:
    header_p = section.header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run("LOCKNOTE  /  SUPABASE LEARNING NOTES")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    label = footer_p.add_run("PERSONAL REFERENCE  |  PAGE ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_number(footer_p)


def add_sources(doc: Document) -> None:
    doc.add_heading("Official references", level=1)
    sources = [
        ("Supabase CLI workflow", "https://supabase.com/docs/guides/local-development/cli-workflows"),
        ("Database migrations", "https://supabase.com/docs/guides/local-development/database-migrations"),
        ("Deploy Edge Functions", "https://supabase.com/docs/guides/functions/deploy"),
        ("Auth redirect URLs", "https://supabase.com/docs/guides/auth/redirect-urls"),
        ("Native mobile deep linking", "https://supabase.com/docs/guides/auth/native-mobile-deep-linking"),
        ("Expo React Native quickstart", "https://supabase.com/docs/guides/getting-started/quickstarts/expo-react-native"),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        apply_num(p, 43)
        add_hyperlink(p, label, url)


def build() -> None:
    doc = Document()
    setup_styles(doc)
    for section in doc.sections:
        configure_section(section)
        add_header_footer(section)

    props = doc.core_properties
    props.title = "Supabase Setup for LockNote"
    props.subject = "Personal learning notes for Supabase authentication, sync, and sharing"
    props.author = "LockNote"
    props.keywords = "Supabase, LockNote, Expo, authentication, database, RLS, sync, sharing"

    add_numbering_definition(doc, 43, 43, "bullet", "•")
    add_numbering_definition(doc, 44, 44, "decimal", "%1.")

    # Opening block: editorial-cover-inspired, compact enough to keep the guide practical.
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(30)
    kicker.paragraph_format.space_after = Pt(8)
    set_run_font(kicker.add_run("PERSONAL LEARNING GUIDE"), size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(7)
    set_keep_with_next(title)
    set_run_font(title.add_run("Supabase Setup for LockNote"), size=29, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(
        subtitle.add_run("A beginner-friendly guide to authentication, database setup, sync, and note sharing."),
        size=13,
        color=MUTED,
    )

    add_callout(
        doc,
        "Core idea",
        "LockNote stays local-first: notes are edited and stored on the device. Supabase adds optional account login, manual multi-device sync, and collaboration when a note is explicitly shared.",
    )

    doc.add_heading("What is Supabase?", level=1)
    doc.add_paragraph(
        "Supabase is a backend platform built around PostgreSQL. It gives an app a hosted database plus services such as user authentication, automatically generated APIs, Realtime updates, file storage, and server-side Edge Functions. It is useful when an app needs cloud features without building every backend component from scratch."
    )
    doc.add_paragraph(
        "For LockNote, Supabase is the cloud layer. The app's SQLite database on Android and iOS, and AsyncStorage on web, remain the normal working storage. A signed-in user chooses when to sync private data, while shared notes use Supabase so invited accounts can open and edit the same saved note."
    )

    doc.add_heading("The Supabase pieces used by LockNote", level=2)
    add_definition(doc, "Auth", "Registers users, confirms email addresses, signs users in, restores sessions, and supports password-reset links.")
    add_definition(doc, "PostgreSQL database", "Stores account-owned folders and notes, plus shared-note membership and saved revisions.")
    add_definition(doc, "Row Level Security (RLS)", "Database rules that limit each signed-in user to rows they own or have permission to access.")
    add_definition(doc, "RPC", "A PostgreSQL function callable through Supabase. LockNote uses sync_private_data for an atomic two-way merge.")
    add_definition(doc, "Edge Function", "Server-side TypeScript code. The share-note function safely looks up a registered account by email without exposing privileged keys to the app.")
    add_definition(doc, "Realtime", "Notifies collaborators when shared-note or membership records change so the local cache can refresh.")

    add_callout(
        doc,
        "Security boundary",
        "RLS protects cloud rows, but LockNote does not end-to-end encrypt note content before upload. A LockNote password is an access gate, not encryption. Do not store the service-role key, database password, or any other secret in the Expo app.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.add_page_break()
    doc.add_heading("1. Know your project values", level=1)
    doc.add_paragraph("Your current Supabase project reference is nhsomigoubtuajiloenc. The app-facing project URL is:")
    add_code(doc, ["https://nhsomigoubtuajiloenc.supabase.co"])

    doc.add_heading("Publishable key versus secrets", level=2)
    doc.add_paragraph(
        "A key beginning with sb_publishable_ is the client-safe publishable key. It fills the same role commonly called the anon key. LockNote keeps the existing environment variable name SUPABASE_ANON_KEY, even when its value uses the newer publishable-key format."
    )
    add_bullet(doc, "Safe in the client bundle: the project URL and publishable key, provided RLS is enabled and correct.")
    add_bullet(doc, "Never put in the app: sb_secret_ keys, service_role keys, the database password, access tokens, or refresh tokens.")
    add_bullet(doc, "Never paste a real key into screenshots, public repositories, support messages, or learning notes.")

    doc.add_heading("2. Configure the LockNote environment", level=1)
    doc.add_paragraph("At the LockNote project root, create or update .env with these two values:")
    add_code(doc, [
        "SUPABASE_URL=https://nhsomigoubtuajiloenc.supabase.co",
        "SUPABASE_ANON_KEY=sb_publishable_your_key_here",
    ])
    add_callout(
        doc,
        "Why the name says ANON_KEY",
        "The source code already reads SUPABASE_ANON_KEY. Keeping that name avoids unnecessary app changes; the value may still be a modern sb_publishable_ key.",
    )
    add_step(doc, "Save the .env file. ", "Do not commit it to Git.")
    add_step(doc, "Restart Expo. ", "Configuration is loaded when Metro starts, so a running development server will not automatically receive the new values.")
    add_code(doc, ["npm.cmd start -- --clear"])

    doc.add_heading("3. Prepare the Supabase CLI", level=1)
    doc.add_paragraph(
        "The LockNote repository already contains SQL migrations and the share-note Edge Function. The safest setup is to deploy those versioned files rather than manually creating tables in the dashboard."
    )
    add_step(doc, "Initialize the local Supabase configuration. ", "Run this once if supabase/config.toml does not exist.")
    add_code(doc, ["npx.cmd supabase@latest init"])
    add_step(doc, "Sign in to the CLI. ", "A browser or access-token flow connects the command line to your Supabase account.")
    add_code(doc, ["npx.cmd supabase@latest login"])
    add_step(doc, "Link this repository to the cloud project. ", "Use the project reference from the dashboard URL.")
    add_code(doc, ["npx.cmd supabase@latest link --project-ref nhsomigoubtuajiloenc"])
    add_callout(
        doc,
        "Password prompt",
        "If the CLI asks for a database password, use the database password from Supabase project settings. The publishable key is not the database password.",
        fill=PALE_GOLD,
        accent=GOLD,
    )

    doc.add_page_break()
    doc.add_heading("4. Review and deploy the database", level=1)
    doc.add_paragraph("Preview the changes first, then apply them:")
    add_code(doc, [
        "npx.cmd supabase@latest db push --dry-run",
        "npx.cmd supabase@latest db push",
    ])
    add_callout(
        doc,
        "Before confirming",
        "Read the dry-run output. A normal first deployment should create the LockNote tables, functions, triggers, RLS policies, and Realtime publication changes. Stop if it proposes unexpected destructive changes.",
    )

    doc.add_heading("Migrations currently in LockNote", level=2)
    for migration in (
        "202608230001_collaboration_release_1.sql - shared-note collaboration",
        "202608240001_private_note_sync.sql - private folder/note account sync",
        "202608250001_note_archive.sql - note archive state",
        "202608250002_folder_archive.sql - folder archive state",
    ):
        add_bullet(doc, migration)

    doc.add_heading("Expected database objects", level=2)
    add_definition(doc, "profiles", "A safe account profile used for collaboration and registered-email lookup.")
    add_definition(doc, "shared_notes", "The latest saved shared-note snapshot, owner, revision, and last-editor metadata.")
    add_definition(doc, "note_members", "Which registered accounts may access a shared note.")
    add_definition(doc, "private_folders", "Folders belonging to one authenticated account for manual sync.")
    add_definition(doc, "private_notes", "Private or owned notes belonging to one authenticated account for manual sync.")
    add_definition(doc, "sync_private_data", "The RPC used to merge local and cloud snapshots with timestamps and deletion tombstones.")

    doc.add_heading("5. Deploy the sharing function", level=1)
    doc.add_paragraph("Deploy the authenticated function that accepts an email address and adds a registered collaborator:")
    add_code(doc, ["npx.cmd supabase@latest functions deploy share-note"])
    add_paragraph = doc.add_paragraph(
        "The service-role key belongs only in Supabase's server-side function environment. LockNote's mobile and web bundles must never contain it."
    )
    add_paragraph.paragraph_format.keep_with_next = True

    doc.add_heading("6. Configure authentication URLs", level=1)
    doc.add_paragraph(
        "In Supabase Dashboard, open Authentication > URL Configuration. Add the following Redirect URLs so email confirmation and both password-recovery paths can return to LockNote:"
    )
    add_code(doc, [
        "locknote://auth-confirm",
        "locknote://reset-password",
        "locknote://reset-lock-password",
        "http://localhost:8081/**",
    ])
    add_bullet(doc, "Keep localhost only for local web development.")
    add_bullet(doc, "Add the deployed web app URL later, for example https://your-domain.example/**.")
    add_bullet(doc, "Use the exact locknote scheme because app.config.js declares scheme: locknote.")

    doc.add_page_break()
    doc.add_heading("7. Verify the backend", level=1)
    doc.add_paragraph("Use this checklist after deployment:")
    add_check(doc, "Table Editor shows profiles, shared_notes, note_members, private_folders, and private_notes.")
    add_check(doc, "Database Functions shows sync_private_data.")
    add_check(doc, "Edge Functions shows share-note as deployed.")
    add_check(doc, "Realtime includes shared_notes and note_members.")
    add_check(doc, "RLS is enabled on every account or collaboration table.")
    add_check(doc, "Authentication > URL Configuration contains all native redirect URLs.")
    add_check(doc, "The Expo server was restarted after changing .env.")

    doc.add_heading("8. Test LockNote end to end", level=1)
    add_step(doc, "Register account A. ", "Confirm its email if email confirmation is enabled, then sign in.")
    add_step(doc, "Create private data. ", "Add a folder and several note types, then run Profile > Sync Notes.")
    add_step(doc, "Test a second device or browser. ", "Sign in as account A, run Sync Notes, and confirm the same folders and notes appear.")
    add_step(doc, "Register account B. ", "The collaborator must have a LockNote account before sharing by email can find it.")
    add_step(doc, "Share from account A. ", "Open a note, invite account B's normalized email address, and check that the share indicator appears on Home.")
    add_step(doc, "Edit as account B. ", "Open the Shared tab, edit the note, and confirm account A receives the saved update and sees the last-editor footer.")
    add_step(doc, "Test recovery. ", "Request account-password and LockNote-password reset emails and confirm each callback opens the correct in-app flow.")

    doc.add_heading("How LockNote sync behaves", level=2)
    add_bullet(doc, "Editing is local and offline-first; manual Sync Notes is the explicit private-data cloud action.")
    add_bullet(doc, "Folders are applied before notes so folder relationships remain valid.")
    add_bullet(doc, "A null folder_id means a note belongs on Home.")
    add_bullet(doc, "Deletion tombstones prevent an older device from restoring deleted content during a later sync.")
    add_bullet(doc, "Shared-with-me notes are excluded from private sync because collaboration is their source of truth.")

    doc.add_page_break()
    doc.add_heading("Troubleshooting", level=1)
    trouble = [
        ("Account services are not configured correctly", "Check SUPABASE_URL and SUPABASE_ANON_KEY, remove accidental quotes or spaces, then restart Expo with --clear."),
        ("The CLI will not link or push", "Confirm you are logged in, the project reference is correct, and any password prompt is answered with the database password rather than the publishable key."),
        ("A confirmation or reset link opens the wrong place", "Add the exact callback to Authentication > URL Configuration and make sure the link uses the locknote scheme on native."),
        ("Sharing cannot find an email address", "The invited person must register first. Confirm the profiles trigger/backfill migration ran and the email is normalized."),
        ("A shared note does not refresh", "Confirm Realtime publication includes shared_notes and note_members, then verify both devices have valid sessions."),
        ("Sync fails with a permission error", "Check that RLS policies were created and the request uses the signed-in user's session, not a missing or expired session."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660], indent_dxa=120)
    set_repeat_table_header(table.rows[0])
    table.rows[0].cells[0].text = "Symptom"
    table.rows[0].cells[1].text = "What to check"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=10, color=NAVY, bold=True)
    for symptom, fix in trouble:
        cells = table.add_row().cells
        cells[0].text = symptom
        cells[1].text = fix
        for idx, cell in enumerate(cells):
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    set_run_font(run, size=9.5, color=NAVY, bold=(idx == 0))

    add_callout(
        doc,
        "Avoid destructive commands",
        "Do not run db reset --linked against a real project unless you deliberately intend to erase and rebuild its remote database.",
        fill=PALE_RED,
        accent=RED,
        ink=RED,
    )

    doc.add_heading("My setup status", level=1)
    add_check(doc, "Environment variables saved locally; no secrets committed.")
    add_check(doc, "Supabase CLI initialized, authenticated, and linked.")
    add_check(doc, "Dry run reviewed and migrations deployed.")
    add_check(doc, "share-note Edge Function deployed.")
    add_check(doc, "Auth redirect URLs configured.")
    add_check(doc, "Registration, confirmation, sign-in, persistence, and recovery tested.")
    add_check(doc, "Private sync tested between two clients.")
    add_check(doc, "Sharing tested with two separate registered accounts.")

    add_sources(doc)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(note.add_run("Prepared as personal learning notes for the LockNote project."), size=9, color=MUTED, italic=True)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
